"""
modules/docx_builder.py — Build .docx files for resume and cover letter export.

Formatting matches the target resume template exactly:
  - Font: Arial throughout
  - Colors: Dark (#1F2D3D), Gray (#555555), Role gray (#444444), Blue (#1F5C8B)
  - Section headers: 11pt bold blue with bottom border (#2E75B6)
  - Company: 11pt bold dark, location tab-separated in gray
  - Role: 10pt gray (#444444)
  - Bullets: filled circle (●), List Paragraph style
  - Core Skills: plain paragraphs with bold label, no bullet
  - Additional Experience: bold label + gray text, 9.5pt
"""
from __future__ import annotations
import io
import logging
import os
import re
from datetime import date

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Regex helpers
# ---------------------------------------------------------------------------

_SECTION_HEADER_RE = re.compile(r"^[A-Z][A-Z0-9\s&/\-]{2,}$")
_BULLET_RE         = re.compile(r"^[•\-]\s+(.+)$")
_COMPANY_RE        = re.compile(r"^.+\s+\|\s+.+$")
_ROLE_RE           = re.compile(r".+[•·]\s*\d{4}")
_INLINE_BOLD_RE    = re.compile(r"^([A-Za-z][A-Za-z0-9\s&/\-]{1,40}):\s+(.+)$")
_ADDITIONAL_EXP_RE = re.compile(r"^Additional Experience:\s*(.+)$", re.IGNORECASE)
_SELECTED_PROJ_RE  = re.compile(r"^Selected Projects$", re.IGNORECASE)

FONT = "Arial"


def _sanitize(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text or "")


def _set_font(run, size_pt, color_rgb=None, bold=None, font_name=FONT):
    """Apply font properties to a run."""
    from docx.shared import Pt, RGBColor
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def _add_bottom_border(paragraph, color: str = "2E75B6", sz: str = "4"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet_numbering(paragraph, doc):
    """Add bullet list numbering to a paragraph, matching the target format."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Ensure numbering part exists and create abstract + concrete num if needed
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part._element

    # Check if we already created our numId
    if not hasattr(doc, '_resume_num_id'):
        # Find highest abstractNumId
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        abstract_nums = numbering_elm.findall('.//w:abstractNum', ns)
        max_abstract = max((int(a.get(qn('w:abstractNumId'))) for a in abstract_nums), default=0)
        new_abstract_id = max_abstract + 1

        # Create abstractNum with filled circle bullet
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(new_abstract_id))
        multi = OxmlElement('w:multiLevelType')
        multi.set(qn('w:val'), 'hybridMultilevel')
        abstract_num.append(multi)

        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'bullet')
        lvl.append(num_fmt)
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), '\u25CF')  # ● filled circle
        lvl.append(lvl_text)
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)
        pPr_lvl = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:hanging'), '360')
        pPr_lvl.append(ind)
        lvl.append(pPr_lvl)
        abstract_num.append(lvl)

        numbering_elm.append(abstract_num)

        # Create num referencing the abstractNum
        nums = numbering_elm.findall('.//w:num', ns)
        max_num = max((int(n.get(qn('w:numId'))) for n in nums), default=0)
        new_num_id = max_num + 1

        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(new_num_id))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(new_abstract_id))
        num.append(abstract_ref)
        numbering_elm.append(num)

        doc._resume_num_id = new_num_id

    # Apply numbering to paragraph
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numPr.append(ilvl)
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(doc._resume_num_id))
    numPr.append(numId)
    pPr.append(numPr)


# Color palette (matching target docx)
_DARK   = (0x1F, 0x2D, 0x3D)  # name, company, body text
_GRAY   = (0x55, 0x55, 0x55)  # contact, location
_BLUE   = (0x1F, 0x5C, 0x8B)  # tagline, section headers
_ROLE   = (0x44, 0x44, 0x44)  # role line


# ---------------------------------------------------------------------------
# Resume
# ---------------------------------------------------------------------------

def build_resume_docx(resume_text: str, template_path: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    resume_text = _sanitize(resume_text)

    if template_path and os.path.isfile(template_path):
        return _inject_into_template(resume_text, template_path, "[RESUME_CONTENT]")

    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10)

    # Margins — match target: ~0.55in top/bottom, ~0.5in left/right
    for sec in doc.sections:
        sec.top_margin    = Emu(502920)
        sec.bottom_margin = Emu(502920)
        sec.left_margin   = Emu(640080)
        sec.right_margin  = Emu(640080)

    lines = [l.rstrip() for l in resume_text.splitlines()]

    # ── Locate header anchor lines ─────────────────────────────────────────
    name_idx = next((i for i, l in enumerate(lines) if l.strip()), None)
    contact_idx = None
    tagline_idx = None

    if name_idx is not None:
        for i in range(name_idx + 1, min(name_idx + 4, len(lines))):
            l = lines[i].strip()
            if l and ("•" in l or "@" in l or re.search(r"\d{3}[-.\s]\d{3}", l)):
                contact_idx = i
                break
        start = (contact_idx + 1) if contact_idx is not None else (name_idx + 1)
        for i in range(start, min(start + 5, len(lines))):
            l = lines[i].strip()
            if l and not _SECTION_HEADER_RE.match(l):
                tagline_idx = i
                break

    # Track which section we're in to handle Core Skills differently
    current_section = None

    # ── Render ──────────────────────────────────────────────────────────────
    for idx, line in enumerate(lines):
        stripped = line.strip()

        # ── Name ──
        if idx == name_idx:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped)
            _set_font(r, 18, _DARK, bold=True)
            continue

        # ── Contact ──
        if idx == contact_idx:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped)
            _set_font(r, 9.5, _GRAY)
            continue

        # ── Tagline ──
        if idx == tagline_idx:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(stripped)
            _set_font(r, 9.5, _BLUE, bold=True)
            continue

        # ── Empty line → spacer ──
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            continue

        # ── Section header (ALL CAPS) ──
        if _SECTION_HEADER_RE.match(stripped):
            current_section = stripped.upper()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(stripped)
            _set_font(r, 11, _BLUE, bold=True)
            _add_bottom_border(p)
            continue

        # ── Additional Experience ──
        m = _ADDITIONAL_EXP_RE.match(stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run("Additional Experience: ")
            _set_font(r1, 9.5, _DARK, bold=True)
            r2 = p.add_run(m.group(1))
            _set_font(r2, 9.5, _GRAY)
            continue

        # ── Selected Projects sub-header ──
        if _SELECTED_PROJ_RE.match(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5.5)
            p.paragraph_format.space_after = Pt(1.8)
            r = p.add_run(stripped)
            _set_font(r, 10, _DARK, bold=True)
            continue

        # ── Bullet line ──
        bm = _BULLET_RE.match(stripped)
        if bm:
            text = bm.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1.6)
            p.paragraph_format.space_after = Pt(1.6)
            _add_bullet_numbering(p, doc)

            # Check for inline bold label (Core Skills or Selected Projects style)
            lm = _INLINE_BOLD_RE.match(text)
            if lm and current_section and "SKILL" in current_section:
                r1 = p.add_run(lm.group(1) + ": ")
                _set_font(r1, 10, _DARK, bold=True)
                r2 = p.add_run(lm.group(2))
                _set_font(r2, 10, _DARK)
            else:
                r = p.add_run(text)
                _set_font(r, 10, _DARK)
            continue

        # ── Role line ("Title • YYYY–YYYY") — check before company ──
        if _ROLE_RE.search(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped)
            _set_font(r, 10, _ROLE)
            continue

        # ── Company line ("Name | Location") ──
        if _COMPANY_RE.match(stripped):
            parts = stripped.split("|", 1)
            company = parts[0].strip()
            location = parts[1].strip() if len(parts) > 1 else ""

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1.5)

            r1 = p.add_run(company)
            _set_font(r1, 11, _DARK, bold=True)
            if location:
                r2 = p.add_run("\t" + location)
                _set_font(r2, 10, _GRAY)
            continue

        # ── Body text (summary, education, etc.) ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(stripped)
        _set_font(r, 10, _DARK)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Cover letter
# ---------------------------------------------------------------------------

def build_cover_letter_docx(cover_letter_text: str, template_path: str = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches

    cover_letter_text = _sanitize(cover_letter_text)

    if template_path and os.path.isfile(template_path):
        return _inject_into_template(cover_letter_text, template_path, "[COVER_LETTER_CONTENT]")

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    p = doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    p.paragraph_format.space_after = Pt(12)
    for r in p.runs:
        r.font.name = FONT
        r.font.size = Pt(11)

    for para in cover_letter_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(10)
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Template injection
# ---------------------------------------------------------------------------

def _inject_into_template(content: str, template_path: str, placeholder: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(template_path)

    target_para = None
    for para in doc.paragraphs:
        if placeholder in para.text:
            target_para = para
            break

    if target_para is None:
        logger.warning(
            "Placeholder '%s' not found in template %s — appending content.",
            placeholder, template_path
        )
        for line in content.splitlines():
            p = doc.add_paragraph(line.rstrip())
            for r in p.runs:
                r.font.size = Pt(10)
    else:
        parent = target_para._element.getparent()
        idx = list(parent).index(target_para._element)
        parent.remove(target_para._element)
        for j, line in enumerate(content.splitlines()):
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = line.rstrip()
            new_r.append(new_t)
            new_p.append(new_r)
            parent.insert(idx + j, new_p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
